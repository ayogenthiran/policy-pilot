"""
Document loading utilities for Policy Pilot RAG backend.
Handles loading and parsing of different document types.
"""

import io
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import PyPDF2
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType

from src.core.logging import get_logger
from src.utils.exceptions import DocumentProcessingError, ValidationError

logger = get_logger(__name__)


class DocumentLoader:
    """Class for loading and parsing different document types."""
    
    def __init__(self):
        """Initialize the document loader."""
        self.supported_types = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,  # python-docx can handle .doc files too
            '.txt': self._load_txt,
            '.md': self._load_txt,
            '.rtf': self._load_txt
        }
    
    def load_document(self, file_path: Union[str, Path], file_type: str) -> Dict[str, Any]:
        """
        Load document based on file type.
        
        Args:
            file_path: Path to the document file
            file_type: File extension (e.g., '.pdf', '.docx')
            
        Returns:
            Dictionary with text_content and metadata
            
        Raises:
            DocumentProcessingError: If document loading fails
            ValidationError: If file type is not supported
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise DocumentProcessingError(f"File not found: {file_path}")
            
            # Normalize file type
            file_type = file_type.lower()
            if not file_type.startswith('.'):
                file_type = f'.{file_type}'
            
            if file_type not in self.supported_types:
                raise ValidationError(f"Unsupported file type: {file_type}")
            
            # Load document
            loader_func = self.supported_types[file_type]
            result = loader_func(file_path)
            
            # Validate result
            if not result.get('text_content'):
                raise DocumentProcessingError("No text content extracted from document")
            
            logger.info(f"Successfully loaded document: {file_path} ({file_type})")
            return result
            
        except (DocumentProcessingError, ValidationError):
            raise
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise DocumentProcessingError(f"Failed to load document: {e}")
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Load PDF document using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with text_content and metadata
        """
        try:
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    pdf_metadata = pdf_reader.metadata
                    metadata.update({
                        'title': pdf_metadata.get('/Title', ''),
                        'author': pdf_metadata.get('/Author', ''),
                        'subject': pdf_metadata.get('/Subject', ''),
                        'creator': pdf_metadata.get('/Creator', ''),
                        'producer': pdf_metadata.get('/Producer', ''),
                        'creation_date': str(pdf_metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_metadata.get('/ModDate', ''))
                    })
                
                # Extract text from all pages
                total_pages = len(pdf_reader.pages)
                metadata['total_pages'] = total_pages
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append({
                                'page_number': page_num,
                                'content': page_text.strip(),
                                'word_count': len(page_text.split()),
                                'char_count': len(page_text)
                            })
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
                
                # Add document-level metadata
                metadata.update({
                    'file_type': 'pdf',
                    'total_pages_processed': len(text_content),
                    'extraction_method': 'PyPDF2'
                })
                
                logger.info(f"PDF loaded: {total_pages} pages, {len(text_content)} pages with content")
                return {
                    'text_content': text_content,
                    'metadata': metadata
                }
                
        except Exception as e:
            logger.error(f"Failed to load PDF {file_path}: {e}")
            raise DocumentProcessingError(f"Failed to load PDF document: {e}")
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Load DOCX document using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with text_content and metadata
        """
        try:
            text_content = []
            metadata = {}
            
            doc = DocxDocument(file_path)
            
            # Extract document properties
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or ''
            })
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'content': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal',
                        'word_count': len(para.text.split()),
                        'char_count': len(para.text)
                    })
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    tables.append({
                        'content': '\n'.join(table_text),
                        'rows': len(table_text),
                        'word_count': len(' '.join(table_text).split()),
                        'char_count': len(' '.join(table_text))
                    })
            
            # Combine paragraphs and tables
            text_content.extend(paragraphs)
            text_content.extend(tables)
            
            # Add document-level metadata
            metadata.update({
                'file_type': 'docx',
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'total_elements': len(text_content),
                'extraction_method': 'python-docx'
            })
            
            logger.info(f"DOCX loaded: {len(paragraphs)} paragraphs, {len(tables)} tables")
            return {
                'text_content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            raise DocumentProcessingError(f"Failed to load DOCX document: {e}")
    
    def _load_txt(self, file_path: Path) -> Dict[str, Any]:
        """
        Load text file (TXT, MD, RTF).
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with text_content and metadata
        """
        try:
            text_content = []
            metadata = {}
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise DocumentProcessingError("Could not decode file with any supported encoding")
            
            # Split content into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            for i, para in enumerate(paragraphs, 1):
                text_content.append({
                    'paragraph_number': i,
                    'content': para,
                    'word_count': len(para.split()),
                    'char_count': len(para)
                })
            
            # Add document-level metadata
            metadata.update({
                'file_type': file_path.suffix.lower(),
                'total_paragraphs': len(paragraphs),
                'total_elements': len(text_content),
                'encoding_used': encoding,
                'extraction_method': 'text_file'
            })
            
            logger.info(f"Text file loaded: {len(paragraphs)} paragraphs")
            return {
                'text_content': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to load text file {file_path}: {e}")
            raise DocumentProcessingError(f"Failed to load text file: {e}")
    
    def get_supported_types(self) -> List[str]:
        """
        Get list of supported file types.
        
        Returns:
            List of supported file extensions
        """
        return list(self.supported_types.keys())
    
    def validate_file_type(self, file_type: str) -> bool:
        """
        Validate if file type is supported.
        
        Args:
            file_type: File extension to validate
            
        Returns:
            True if supported, False otherwise
        """
        file_type = file_type.lower()
        if not file_type.startswith('.'):
            file_type = f'.{file_type}'
        return file_type in self.supported_types
